#!/usr/bin/env python3
"""
Weekly Meeting List Generator
Reads an Outlook calendar PDF, lets you select and categorise meetings
via a browser UI, then generates a colour-coded Word document.
"""

import argparse
import json
import os
import re
import sys
import threading
import webbrowser
from datetime import date
from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

COLORS = {
    "Council":    {"fill": "DDEEFF", "text": "0C447C"},
    "Parliament": {"fill": "FDDEDE", "text": "A32D2D"},
    "Internal":   {"fill": "E8F5E0", "text": "3B6D11"},
}

DAY_HEADER_COLOR = "4472C4"
COL_HEADER_COLOR = "2E5FAB"

PAGE_WIDTH_CM = 21.0
MARGIN_CM     = 1.9
CONTENT_CM    = PAGE_WIDTH_CM - 2 * MARGIN_CM  # 17.2 cm
COL_TIME_CM   = 2.5
COL_PART_CM   = 4.5
COL_TITLE_CM  = CONTENT_CM - COL_TIME_CM - COL_PART_CM  # 10.2 cm

DAY_ORDER = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]


# ---------------------------------------------------------------------------
# PDF Parsing
# ---------------------------------------------------------------------------

DAY_RE  = re.compile(r"^(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),\s+\d{1,2}\s+\w+", re.IGNORECASE)
TIME_RE = re.compile(r"^\d{1,2}:\d{2}\s*[-–]\s*\d{1,2}:\d{2}")
NOISE_RE = re.compile(r"^(08:00\s*PRIVATE|PRIVATE|https?://\S+)$", re.IGNORECASE)


def parse_column(lines: list[str]) -> list[dict]:
    """Parse a list of text lines from one column into a list of meeting dicts."""
    meetings = []
    current_day   = None
    current_time  = None
    current_parts = []

    def flush():
        nonlocal current_time, current_parts
        if current_day and current_time and current_parts:
            title = " ".join(current_parts).strip()
            title = re.sub(r"\s*-\s*$", "", title).strip()
            if title:
                meetings.append({"day": current_day, "time": current_time, "title": title})
        current_time  = None
        current_parts = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if DAY_RE.match(line):
            flush()
            current_day = line
            continue

        if current_day is None:
            continue

        if NOISE_RE.match(line):
            flush()
            continue

        m = TIME_RE.match(line)
        if m:
            flush()
            time_str  = m.group(0).replace(" ", "").replace("-", "–")
            remainder = re.sub(r"\s*-\s*$", "", line[m.end():].strip()).strip()
            current_time  = time_str
            current_parts = [remainder] if remainder else []
            continue

        if current_time is not None:
            cleaned = re.sub(r"\s*-\s*$", "", line).strip()
            if cleaned:
                current_parts.append(cleaned)

    flush()
    return meetings


def parse_pdf(pdf_path: str) -> list[dict]:
    """
    Extract meetings from an Outlook weekly PDF.
    Outlook lays out Mon+Tue, Wed+Thu, Fri in two-column format.
    We split each page into left and right halves and parse each independently.
    """
    all_entries = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(keep_blank_chars=False, x_tolerance=3, y_tolerance=3)
            if not words:
                continue

            page_mid = page.width / 2

            # Group words into lines by y-position
            lines_by_y: dict[int, list] = {}
            for w in words:
                y_key = round(w["top"] / 4) * 4
                lines_by_y.setdefault(y_key, []).append(w)

            sorted_ys = sorted(lines_by_y.keys())

            # Build left and right text line lists independently
            left_lines  = []
            right_lines = []

            for y in sorted_ys:
                row_words = lines_by_y[y]
                left_words  = sorted([w for w in row_words if w["x0"] <  page_mid], key=lambda w: w["x0"])
                right_words = sorted([w for w in row_words if w["x0"] >= page_mid], key=lambda w: w["x0"])
                if left_words:
                    left_lines.append(" ".join(w["text"] for w in left_words))
                if right_words:
                    right_lines.append(" ".join(w["text"] for w in right_words))

            all_entries.extend(parse_column(left_lines))
            all_entries.extend(parse_column(right_lines))

    # Deduplicate
    seen   = set()
    unique = []
    for e in all_entries:
        key = (e["day"], e["time"], e["title"])
        if key not in seen:
            seen.add(key)
            unique.append(e)

    # Sort by day order then start time
    def sort_key(e):
        day_name = e["day"].split(",")[0].strip().capitalize()
        d = DAY_ORDER.index(day_name) if day_name in DAY_ORDER else 99
        t = e["time"].replace("–", "-").split("-")[0].strip().replace(":", "")
        return (d, t.zfill(4))

    unique.sort(key=sort_key)
    return unique


# ---------------------------------------------------------------------------
# HTML UI
# ---------------------------------------------------------------------------

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Weekly Meeting List</title>
<style>
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: Arial, sans-serif; font-size: 14px; background: #f5f5f5; color: #222; padding: 2rem; }
  h1 { font-size: 22px; font-weight: 600; color: #2E5FAB; margin-bottom: 0.25rem; }
  .subtitle { font-size: 13px; color: #666; font-style: italic; margin-bottom: 1.5rem; }
  .day-block { margin-bottom: 1.5rem; border-radius: 6px; overflow: hidden; border: 1px solid #ddd; }
  .day-header { background: #4472C4; color: #fff; font-weight: 600; font-size: 15px; padding: 8px 14px; }
  .meeting-row { display: grid; grid-template-columns: 36px 110px 1fr 160px;
                 align-items: center; border-bottom: 1px solid #eee; background: #fff; }
  .meeting-row:last-child { border-bottom: none; }
  .meeting-row.selected-Council    { background: #DDEEFF; }
  .meeting-row.selected-Parliament { background: #FDDEDE; }
  .meeting-row.selected-Internal   { background: #E8F5E0; }
  .cell-check { display: flex; align-items: center; justify-content: center; padding: 10px 4px; }
  .cell-check input { width: 16px; height: 16px; cursor: pointer; }
  .cell { padding: 10px 12px; }
  .time  { font-size: 12px; color: #555; white-space: nowrap; }
  .title { font-size: 13px; font-weight: 500; line-height: 1.4; }
  .type-select { font-size: 13px; padding: 5px 8px; border: 1px solid #ccc;
                 border-radius: 4px; background: #fff; cursor: pointer; width: 100%; }
  .type-select:disabled { opacity: 0.35; cursor: default; }
  .summary { display: flex; align-items: center; gap: 1rem; margin: 1.5rem 0 1rem;
             padding: 12px 16px; background: #fff; border: 1px solid #ddd;
             border-radius: 6px; flex-wrap: wrap; }
  .summary span { font-size: 13px; color: #555; }
  .summary strong { color: #222; }
  .generate-btn { margin-left: auto; padding: 9px 22px; font-size: 14px; font-weight: 600;
                  background: #2E5FAB; color: #fff; border: none; border-radius: 6px; cursor: pointer; }
  .generate-btn:disabled { background: #aaa; cursor: default; }
  .generate-btn:not(:disabled):hover { background: #1e4a8a; }
  .legend { display: flex; gap: 10px; margin-top: 1.5rem; }
  .legend-item { padding: 5px 16px; border-radius: 4px; font-size: 12px; font-weight: 600; }
  .done-msg { display: none; background: #E8F5E0; border: 1px solid #3B6D11; color: #3B6D11;
              padding: 14px 20px; border-radius: 6px; font-size: 14px; margin-top: 1rem; }
</style>
</head>
<body>
<h1>Weekly Meeting List</h1>
<p class="subtitle">Select the meetings to include and assign each a type. Then click Generate.</p>
<div id="app"></div>
<div class="summary">
  <span><strong id="sel-count">0</strong> meetings selected</span>
  <span><strong id="typed-count">0</strong> typed</span>
  <button class="generate-btn" id="gen-btn" disabled onclick="generate()">Generate Word document</button>
</div>
<div class="legend">
  <div class="legend-item" style="background:#DDEEFF;color:#0C447C;">Council</div>
  <div class="legend-item" style="background:#FDDEDE;color:#A32D2D;">Parliament</div>
  <div class="legend-item" style="background:#E8F5E0;color:#3B6D11;">Internal</div>
</div>
<div class="done-msg" id="done-msg">
  ✓ Word document generated successfully! Check the folder where your PDF is saved.
</div>
<script>
const meetings = MEETINGS_JSON;
const state = {};
meetings.forEach(m => state[m.id] = { checked: false, type: "" });
const days = [...new Set(meetings.map(m => m.day))];

function render() {
  const app = document.getElementById("app");
  let html = "";
  days.forEach(day => {
    const dm = meetings.filter(m => m.day === day);
    html += `<div class="day-block"><div class="day-header">${day}</div>`;
    dm.forEach(m => {
      const s = state[m.id];
      const cls = s.checked && s.type ? ` selected-${s.type}` : "";
      html += `<div class="meeting-row${cls}" id="row-${m.id}">
        <div class="cell-check"><input type="checkbox" ${s.checked?"checked":""} onchange="toggle(${m.id})"></div>
        <div class="cell"><div class="time">${m.time}</div></div>
        <div class="cell"><div class="title">${m.title}</div></div>
        <div class="cell">
          <select class="type-select" id="sel-${m.id}" ${s.checked?"":"disabled"} onchange="setType(${m.id},this.value)">
            <option value="">— type —</option>
            <option value="Council"    ${s.type==="Council"   ?"selected":""}>Council</option>
            <option value="Parliament" ${s.type==="Parliament"?"selected":""}>Parliament</option>
            <option value="Internal"   ${s.type==="Internal"  ?"selected":""}>Internal</option>
          </select>
        </div>
      </div>`;
    });
    html += `</div>`;
  });
  app.innerHTML = html;
  updateSummary();
}

function toggle(id) {
  state[id].checked = !state[id].checked;
  if (!state[id].checked) state[id].type = "";
  render();
}

function setType(id, val) {
  state[id].type = val;
  render();
}

function updateSummary() {
  const sel = meetings.filter(m => state[m.id].checked);
  document.getElementById("sel-count").textContent = sel.length;
  const typed = sel.filter(m => state[m.id].type).length;
  document.getElementById("typed-count").textContent = typed;
  document.getElementById("gen-btn").disabled = sel.length === 0 || typed < sel.length;
}

function generate() {
  const sel = meetings.filter(m => state[m.id].checked);
  const payload = sel.map(m => ({ day: m.day, time: m.time, title: m.title, type: state[m.id].type }));
  fetch("/generate", {
    method: "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify(payload)
  }).then(r => r.json()).then(data => {
    if (data.ok) {
      document.getElementById("done-msg").style.display = "block";
      document.getElementById("gen-btn").disabled = true;
      document.getElementById("gen-btn").textContent = "✓ Generated";
    } else {
      alert("Error: " + data.error);
    }
  });
}

render();
</script>
</body>
</html>
"""


# ---------------------------------------------------------------------------
# Word Document Helpers
# ---------------------------------------------------------------------------

def hex_to_rgb(h: str) -> tuple:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def set_cell_borders(cell, color: str = "CCCCCC"):
    tcPr    = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_width(cell, width_cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_cm / 2.54 * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", 60), ("bottom", 60), ("left", 100), ("right", 100)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def fill_cell(cell, text: str, width_cm: float, fill_color: str,
              text_color: str = "000000", bold: bool = False, font_size: int = 10):
    cell.text = ""
    set_cell_shading(cell, fill_color)
    set_cell_borders(cell)
    set_cell_width(cell, width_cm)
    set_cell_margins(cell)
    p   = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name      = "Arial"
    run.font.size      = Pt(font_size)
    run.font.bold      = bold
    run.font.color.rgb = RGBColor(*hex_to_rgb(text_color))


# ---------------------------------------------------------------------------
# Word Document Generation
# ---------------------------------------------------------------------------

def generate_docx(selected: list[dict], pdf_path: str) -> str:
    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(MARGIN_CM)
    section.right_margin  = Cm(MARGIN_CM)
    section.top_margin    = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)

    all_days   = list(dict.fromkeys(m["day"] for m in selected))
    week_label = f"{all_days[0]} – {all_days[-1]}" if len(all_days) > 1 else all_days[0]

    # Title
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Weekly Meeting List — {week_label}")
    run.font.name = "Arial"; run.font.size = Pt(16); run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb("2E5FAB"))

    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("For review and assignment by team leadership")
    run.font.name = "Arial"; run.font.size = Pt(10); run.font.italic = True
    run.font.color.rgb = RGBColor(*hex_to_rgb("666666"))

    # Main table
    table = doc.add_table(rows=0, cols=3)
    tblPr = table._tbl.tblPr
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(CONTENT_CM / 2.54 * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Column header
    hdr = table.add_row()
    fill_cell(hdr.cells[0], "Time",         COL_TIME_CM,  COL_HEADER_COLOR, "FFFFFF", bold=True)
    fill_cell(hdr.cells[1], "Meeting",      COL_TITLE_CM, COL_HEADER_COLOR, "FFFFFF", bold=True)
    fill_cell(hdr.cells[2], "Participants", COL_PART_CM,  COL_HEADER_COLOR, "FFFFFF", bold=True)

    # Data rows grouped by day
    days = list(dict.fromkeys(m["day"] for m in selected))
    for day in days:
        # Day header row (merged)
        day_row = table.add_row()
        for c in day_row.cells:
            set_cell_shading(c, DAY_HEADER_COLOR)
            set_cell_borders(c, DAY_HEADER_COLOR)
        day_row.cells[0].merge(day_row.cells[1])
        day_row.cells[0].merge(day_row.cells[2])
        merged = day_row.cells[0]
        set_cell_margins(merged)
        merged.paragraphs[0].clear()
        p   = merged.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(day)
        run.font.name = "Arial"; run.font.size = Pt(11); run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

        for m in (x for x in selected if x["day"] == day):
            c   = COLORS[m["type"]]
            row = table.add_row()
            fill_cell(row.cells[0], m["time"],  COL_TIME_CM,  c["fill"])
            fill_cell(row.cells[1], m["title"], COL_TITLE_CM, c["fill"])
            fill_cell(row.cells[2], "",         COL_PART_CM,  c["fill"])

    # Legend
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("Colour legend")
    run.font.name = "Arial"; run.font.size = Pt(10); run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb("444444"))

    leg_table = doc.add_table(rows=1, cols=4)
    leg_row   = leg_table.rows[0]
    items     = [
        ("Council",    COLORS["Council"]["fill"],    COLORS["Council"]["text"]),
        ("Parliament", COLORS["Parliament"]["fill"], COLORS["Parliament"]["text"]),
        ("Internal",   COLORS["Internal"]["fill"],   COLORS["Internal"]["text"]),
    ]
    w_each = 3.5
    for i, (label, fill, tcol) in enumerate(items):
        cell = leg_row.cells[i]
        cell.text = ""
        set_cell_shading(cell, fill)
        set_cell_width(cell, w_each)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.name = "Arial"; run.font.size = Pt(9); run.font.bold = True
        run.font.color.rgb = RGBColor(*hex_to_rgb(tcol))
    rem = leg_row.cells[3]
    rem.text = ""
    set_cell_width(rem, CONTENT_CM - w_each * 3)

    # Generation date
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run(f"Generated on {date.today().strftime('%d %B %Y')}")
    run.font.name = "Arial"; run.font.size = Pt(9); run.font.italic = True
    run.font.color.rgb = RGBColor(*hex_to_rgb("999999"))

    out_path = Path(pdf_path).parent / (Path(pdf_path).stem + "_meeting_list.docx")
    doc.save(str(out_path))
    return str(out_path)


# ---------------------------------------------------------------------------
# Local HTTP Server
# ---------------------------------------------------------------------------

class Handler(BaseHTTPRequestHandler):
    pdf_path:       str             = ""
    meetings:       list            = []
    shutdown_event: threading.Event = None

    def log_message(self, format, *args):
        pass

    def do_GET(self):
        if self.path == "/":
            meetings_json = json.dumps([
                {"id": i, "day": m["day"], "time": m["time"], "title": m["title"]}
                for i, m in enumerate(self.meetings)
            ])
            html = HTML_TEMPLATE.replace("MEETINGS_JSON", meetings_json)
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.end_headers()
            self.wfile.write(html.encode("utf-8"))
        else:
            self.send_response(404)
            self.end_headers()

    def do_POST(self):
        if self.path == "/generate":
            length   = int(self.headers.get("Content-Length", 0))
            body     = self.rfile.read(length)
            selected = json.loads(body)
            try:
                out_path = generate_docx(selected, self.pdf_path)
                response = json.dumps({"ok": True, "path": out_path})
                print(f"\n✓ Document saved to: {out_path}")
            except Exception as e:
                import traceback; traceback.print_exc()
                response = json.dumps({"ok": False, "error": str(e)})
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(response.encode("utf-8"))
            threading.Timer(1.0, self.shutdown_event.set).start()
        else:
            self.send_response(404)
            self.end_headers()


# ---------------------------------------------------------------------------
# Entry Point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate a weekly meeting list Word document from an Outlook PDF export."
    )
    parser.add_argument("pdf", help="Path to the Outlook calendar PDF export")
    args = parser.parse_args()

    pdf_path = os.path.abspath(args.pdf)
    if not os.path.isfile(pdf_path):
        print(f"Error: file not found: {pdf_path}")
        sys.exit(1)

    print(f"Reading PDF: {pdf_path}")
    meetings = parse_pdf(pdf_path)
    if not meetings:
        print("No meetings found in the PDF. Please check the file.")
        sys.exit(1)
    print(f"Found {len(meetings)} meetings across {len(set(m['day'] for m in meetings))} days.")

    shutdown_event         = threading.Event()
    Handler.pdf_path       = pdf_path
    Handler.meetings       = meetings
    Handler.shutdown_event = shutdown_event

    server = HTTPServer(("127.0.0.1", 8765), Handler)
    print("Opening browser — select your meetings and click Generate.")
    threading.Timer(0.5, lambda: webbrowser.open("http://127.0.0.1:8765")).start()

    while not shutdown_event.is_set():
        server.handle_request()

    server.server_close()
    print("Done. You can close the browser tab.")


if __name__ == "__main__":
    main()
